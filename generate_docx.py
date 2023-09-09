from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement

def generate_doc(params_list: list, correct_permutation_params: list, name: str) -> None:
    """
    generate docs documentation
    :param params_list: list with initial data
    :param correct_permutation_params: list with transformed data
    """
    doc = Document()
    n = len(params_list)
    print(correct_permutation_params)

    title = "Расчетно-пояснительная записка"
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    title1 = "К схеме размещения и крепления"
    title1_paragraph = doc.add_paragraph()
    title1_run = title1_paragraph.add_run(title1)
    title1_run.bold = True
    title1_run.font.size = Pt(20)
    title1_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    name_paragraph = doc.add_paragraph()
    name_run = name_paragraph.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    harr_paragraph = doc.add_paragraph()
    harr_run = harr_paragraph.add_run("Характеристика платформы и грузовых мест")
    harr_run.bold = True
    harr_run.font.size = Pt(20)
    harr_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    harr4_paragraph = doc.add_paragraph()
    harr4_run = harr4_paragraph.add_run("Характеристика 4-х осной ж/д платформы")
    harr4_run.bold = True

    harr4_data = doc.add_paragraph()
    harr4_data_run = harr4_data.add_run("Длина пола\t\t13400 мм\n")
    harr4_data_run = harr4_data.add_run("Ширина пола\t\t2870 мм\n")
    harr4_data_run = harr4_data.add_run("Масса тары\t\t21 т\n")
    harr4_data_run = harr4_data.add_run("Высота пола от УГР\t\t1310 мм\n")
    harr4_data_run = harr4_data.add_run("Высота пола от УГР\t\t800 мм\n")
    harr4_data_run = harr4_data.add_run("База платформы\t\t9720 мм")

    harr4_data.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    harrgr_paragraph = doc.add_paragraph()
    harrgr_run = harrgr_paragraph.add_run("Характеристика груза:")
    harrgr_run.bold = True

    table = doc.add_table(rows=n + 1, cols=8)
    table.cell(0, 0).text = '№ п/п'
    table.cell(0, 1).text = '№ груза'
    table.cell(0, 2).text = 'Длина'
    table.cell(0, 3).text = 'Ширина'
    table.cell(0, 4).text = 'Высота'
    table.cell(0, 5).text = 'Кол-во (шт)'
    table.cell(0, 6).text = 'Вес 1 ед (кг)'
    table.cell(0, 7).text = 'Общий вес (кг)'

    for i in range(1, n + 1):
        table.cell(i, 0).text = str(i)
        table.cell(i, 1).text = str(i)
        table.cell(i, 2).text = str(params_list[i - 1][0])
        table.cell(i, 3).text = str(params_list[i - 1][1])
        table.cell(i, 4).text = str(params_list[i - 1][2])
        table.cell(i, 5).text = str(params_list[i - 1][3])
        table.cell(i, 6).text = str(params_list[i - 1][4])
        table.cell(i, 7).text = str(params_list[i - 1][3] * params_list[i - 1][4])

    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table.style = 'Table Grid'

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    date = "09.09.2023"
    date_paragraph = footer.add_paragraph("Дата: ")
    date_run = date_paragraph.add_run(date)
    date_run.bold = True
    date_run.font.size = Pt(10)

    city = "г. Москва"
    city_paragraph = footer.add_paragraph(city)
    city_paragraph.runs[0].bold = True
    city_paragraph.runs[0].font.size = Pt(10)

    doc.add_page_break()

    rass_paragraph = doc.add_paragraph()
    rass_run = rass_paragraph.add_run("РАСЧЕТ ПОЛОЖЕНИЯ ОБЩЕГО ЦЕНТРА ТЯЖЕСТИ ГРУЗА")
    rass_run.bold = True
    rass_run.font.size = Pt(20)
    rass_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    smesh_paragraph = doc.add_paragraph()
    smesh_run = smesh_paragraph.add_run("1. Смещение ЦТ грузов в вагоне")
    smesh_run.bold = True

    dl_paragraph = doc.add_paragraph()
    dl_run = dl_paragraph.add_run("Продольное смещение грузов в вагоне: " + str(correct_permutation_params[0]) + "\n")
    dl_run = dl_paragraph.add_run("Продольное смещение грузов с вагоном: " + str(correct_permutation_params[1]))

    smesh1_paragraph = doc.add_paragraph()
    smesh1_run = smesh1_paragraph.add_run("2. Устойчивость грузов с вагоном")
    smesh1_run.bold = True

    h_paragraph = doc.add_paragraph()
    h_run = h_paragraph.add_run("Высота ЦТ грузов в вагоне: " + str(correct_permutation_params[2]) + "\n")
    h_run = h_paragraph.add_run("Общая высота ЦТ: " + str(correct_permutation_params[3]) + "\n")
    h_run = h_paragraph.add_run("Расчет наветренной поверхности: " + str(correct_permutation_params[4]))

    sili_paragraph = doc.add_paragraph()
    sili_run = sili_paragraph.add_run("3. Расчет сил и усточивости")
    sili_run.bold = True

    m = len(correct_permutation_params[5])

    for i in range(m):
        mn_r = doc.add_paragraph()
        mn_run = mn_r.add_run(f"Расчет сил, действующих на Груз № {i + 1}\nУстойчивость груза № {i + 1} в вагоне")
        mn_run.bold = True

        mn_r1 = doc.add_paragraph()
        mn_run1 = mn_r1.add_run('Продольная инерционная сила: ' + str(correct_permutation_params[5][i][0]) + '\n')
        mn_run1 = mn_r1.add_run('Поперечная инерционная сила: ' + str(correct_permutation_params[5][i][1]) + '\n')
        mn_run1 = mn_r1.add_run('Вертикальная инерционная сила: ' + str(correct_permutation_params[5][i][2]) + '\n')
        mn_run1 = mn_r1.add_run('Ветровая нагрузка: ' + str(correct_permutation_params[5][i][3]) + '\n')
        mn_run1 = mn_r1.add_run('Сила трения в продольном направлении: ' + str(correct_permutation_params[5][i][4]) + '\n')
        mn_run1 = mn_r1.add_run('Сила трения в поперечном направлении: ' + str(correct_permutation_params[5][i][5]) + '\n')
        mn_run1 = mn_r1.add_run('Усилия которые должны восприниматься средствами крепления\n')
        mn_run1 = mn_r1.add_run('Продольное: ' + str(correct_permutation_params[5][i][6]) + '\n')
        mn_run1 = mn_r1.add_run('Поперечное: ' + str(correct_permutation_params[5][i][7]) + '\n')
        mn_run1 = mn_r1.add_run('Коэффициент запаса устойчивости от опрокидывания вдоль вагона: ' + str(correct_permutation_params[5][i][8]) + '\n')
        mn_run1 = mn_r1.add_run('Коэффициент запаса устойчивости от опрокидывания поперек вагона: ' + str(correct_permutation_params[5][i][9]))

    doc.save('Расчетно-пояснительная записка.docx')
